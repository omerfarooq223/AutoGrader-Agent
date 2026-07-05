"""
Extraction utilities for reading student submissions.
Supports: PDF, DOCX, .py, .cpp, .ipynb
"""

import json
import logging
import re
import shutil
import subprocess
import zipfile
import os
import tempfile
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor

import fitz  # PyMuPDF
from docx import Document
import openpyxl

import config

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".py", ".cpp", ".ipynb", ".md", ".txt"}
NESTED_ARCHIVE_EXTENSIONS = {".zip", ".rar", ".7z"}
MAX_FILE_SIZE_MB = 20  # Files larger than this are skipped to avoid huge prompt payloads

_VISION_PROMPT = (
    "Describe what is shown in this image in the context of a student assignment. "
    "Be specific about any code output, charts, diagrams, or handwritten content you see."
)


def _describe_image(image_bytes: bytes) -> str | None:
    """Send an image to Gemini's vision model to bypass strict Groq quotas."""
    try:
        import os
        from google import genai
        from google.genai import types
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            return None
        os.environ.pop("GOOGLE_API_KEY", None)
        client = genai.Client(api_key=api_key)
        model_name = os.environ.get("GEMINI_MODEL", "gemini-flash-latest")

        def _call():
            return client.models.generate_content(
                model=model_name,
                contents=[
                    _VISION_PROMPT,
                    types.Part.from_bytes(data=image_bytes, mime_type="image/png"),
                ]
            )

        # Image description is optional — fail fast, never retry.
        # Add explicit timeout to avoid hanging extraction on one image.
        pool = ThreadPoolExecutor(max_workers=1)
        future = pool.submit(_call)
        try:
            response = future.result(timeout=10)
        finally:
            pool.shutdown(wait=False, cancel_futures=True)
        return response.text.strip()
    except Exception as e:
        logger.debug(f"Vision API call failed for an image: {e}", exc_info=True)
        return None



def extract_zip(zip_path: str, extract_to: str | None = None) -> str:
    """Extract a ZIP file and return the path to the extraction directory."""
    if extract_to is None:
        extract_to = tempfile.mkdtemp(prefix="submissions_")

    with zipfile.ZipFile(zip_path, "r") as zf:
        _validate_archive_members(zf.namelist(), zip_path)
        zf.extractall(extract_to)

    return extract_to


def _validate_archive_members(members: list[str], archive_path: str) -> None:
    """Reject archive members that could escape the target extraction directory."""
    for member in members:
        if not member:
            continue
        member_path = os.path.normpath(member)
        if (
            member_path.startswith("..")
            or os.path.isabs(member_path)
            or member_path == ".."
        ):
            raise ValueError(f"Unsafe path in archive {archive_path}: {member}")


def _extract_with_bsdtar(archive_path: str, extract_to: str) -> bool:
    """Extract RAR/7z using bsdtar/libarchive when available."""
    tool = shutil.which("bsdtar")
    if not tool:
        return False

    listed = subprocess.run(
        [tool, "-tf", archive_path],
        capture_output=True,
        text=True,
        timeout=30,
        check=True,
    )
    members = [line.strip() for line in listed.stdout.splitlines() if line.strip()]
    _validate_archive_members(members, archive_path)

    os.makedirs(extract_to, exist_ok=True)
    subprocess.run(
        [tool, "-xf", archive_path, "-C", extract_to],
        capture_output=True,
        text=True,
        timeout=120,
        check=True,
    )
    return True


def _extract_with_7z(archive_path: str, extract_to: str) -> bool:
    """Extract RAR/7z using 7z when available."""
    tool = shutil.which("7z")
    if not tool:
        return False

    listed = subprocess.run(
        [tool, "l", "-slt", archive_path],
        capture_output=True,
        text=True,
        timeout=30,
        check=True,
    )
    members = [
        line.split("=", 1)[1].strip()
        for line in listed.stdout.splitlines()
        if line.startswith("Path = ") and line.split("=", 1)[1].strip() != archive_path
    ]
    _validate_archive_members(members, archive_path)

    os.makedirs(extract_to, exist_ok=True)
    subprocess.run(
        [tool, "x", "-y", f"-o{extract_to}", archive_path],
        capture_output=True,
        text=True,
        timeout=120,
        check=True,
    )
    return True


def extract_archive(archive_path: str, extract_to: str | None = None) -> str:
    """
    Extract a supported archive and return the extraction directory.
    ZIP uses Python's stdlib. RAR/7z require a local tool such as bsdtar or 7z.
    """
    if extract_to is None:
        extract_to = tempfile.mkdtemp(prefix="submissions_")

    ext = Path(archive_path).suffix.lower()
    if ext == ".zip":
        return extract_zip(archive_path, extract_to)

    if ext in {".rar", ".7z"}:
        errors: list[str] = []
        for extractor in (_extract_with_bsdtar, _extract_with_7z):
            try:
                if extractor(archive_path, extract_to):
                    return extract_to
            except Exception as exc:
                errors.append(str(exc))
                logger.debug("Archive extractor failed for %s: %s", archive_path, exc, exc_info=True)
        detail = f" Last error: {errors[-1]}" if errors else ""
        raise RuntimeError(
            "RAR/7z extraction requires a local archive tool such as bsdtar/libarchive or 7z."
            + detail
        )

    raise ValueError(f"Unsupported archive format: {ext}")


def read_pdf(file_path: str) -> str:
    """Extract text and embedded images from a PDF file using PyMuPDF."""
    text_parts: list[str] = []
    with fitz.open(file_path) as doc:
        for page in doc:
            page_text = page.get_text() or ""

            if config.EXTRACT_IMAGES:
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    try:
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        description = _describe_image(image_bytes)
                        if description:
                            page_text += f"\n[Image: {description}]"
                    except Exception:
                        logger.debug("Failed to extract PDF image xref=%s, skipping.", xref, exc_info=True)

            if page_text.strip():
                text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def read_docx(file_path: str) -> str:
    """Extract text and embedded images from a DOCX file using python-docx."""
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Extract table content — tables are not in doc.paragraphs
    table_parts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                table_parts.append(row_text)

    all_parts = paragraphs + (["[Table Content]"] + table_parts if table_parts else [])
    text = "\n".join(all_parts)

    if config.EXTRACT_IMAGES:
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_bytes = rel.target_part.blob
                        description = _describe_image(image_bytes)
                        if description:
                            text += f"\n[Image: {description}]"
                    except Exception:
                        logger.debug("Failed to extract DOCX image, skipping.", exc_info=True)
        except Exception:
            logger.debug("Failed to iterate DOCX relationships, skipping images.", exc_info=True)

    return text.strip()


def read_text_file(file_path: str) -> str:
    """Read plain-text source files (.py, .cpp)."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read().strip()


def read_notebook(file_path: str) -> str:
    """Extract code and markdown cell contents from a Jupyter notebook (.ipynb)."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        notebook = json.load(f)

    cells = notebook.get("cells", [])
    parts: list[str] = []

    for cell in cells:
        cell_type = cell.get("cell_type", "")
        source_lines = cell.get("source", [])
        source = "".join(source_lines).strip()

        if not source:
            continue

        if cell_type == "markdown":
            parts.append(f"[Markdown]\n{source}")
        elif cell_type == "code":
            cell_part = f"[Code]\n{source}"
            # Include cell outputs — printed results, errors, and return values
            outputs = cell.get("outputs", [])
            output_texts = []
            for output in outputs:
                # Stream output (print statements)
                if output.get("output_type") in ("stream", "display_data"):
                    output_texts.append("".join(output.get("text", [])))
                # Execution result (last expression value)
                elif output.get("output_type") == "execute_result":
                    output_texts.append("".join(output.get("text", [])))
                # Errors and tracebacks
                elif output.get("output_type") == "error":
                    output_texts.append(
                        f"ERROR: {output.get('ename')}: {output.get('evalue')}"
                    )
            if output_texts:
                cell_part += f"\n[Output]\n{''.join(output_texts)}"
            parts.append(cell_part)

    return "\n\n".join(parts).strip()


_READERS = {
    ".pdf":   read_pdf,
    ".docx":  read_docx,
    ".py":    read_text_file,
    ".cpp":   read_text_file,
    ".ipynb": read_notebook,
    ".md":    read_text_file,
    ".txt":   read_text_file,
}


# Improved regex patterns for ID and name extraction
# ID patterns: student ID, roll number, registration, etc.
_ID_TOKEN_RE = re.compile(
    # Student ID patterns: CS2024001, F2023376425, B123456, etc.
    r"\b([A-Z]{1,3}\d{5,}|\d{6,}|[A-Z]\d{5,}|[A-Z]{2}\d{4,})\b",
    re.IGNORECASE
)
_NAME_LABEL_RE = re.compile(
    r"(?im)^\s*(?:student\s*)?name\s*[:#\-=]\s*([A-Za-z][A-Za-z .'-]{2,100})\s*$"
)
_ID_LABEL_RE = re.compile(
    r"(?im)^\s*(?:student\s*)?(?:id|roll\s*no|roll#|registration\s*no|reg\s*no|matric|student\s*id)\s*[:#\-=]\s*([A-Za-z0-9\-_/]{4,})"
)

_NOISE_TOKENS = {
    "assignment", "submission", "assignsubmission", "file", "student", "answer",
    "solution", "final", "draft", "copy", "doc", "pdf", "python", "cpp",
    "ipynb", "name", "id", "roll", "registration", "reg", "no", "class",
    "section", "semester", "course", "project", "report", "lab",
}


def _normalize_name(name: str) -> str:
    """Normalize and validate student name."""
    # Remove excessive whitespace
    cleaned = re.sub(r"\s+", " ", name).strip(" _-.,")
    if not cleaned:
        return ""
    
    # Reject if too many numbers (likely not a real name)
    num_count = sum(1 for c in cleaned if c.isdigit())
    if num_count > len(cleaned) * 0.3:  # More than 30% digits = likely not a name
        return ""
    
    # Reject if too short
    if len(cleaned) < 3:
        return ""
    
    # Capitalize words properly
    words = []
    for word in cleaned.split(" "):
        if not word or len(word) < 2:
            continue
        if "'" in word:
            words.append("'".join(part.capitalize() for part in word.split("'")))
        else:
            words.append(word.capitalize())
    
    result = " ".join(words)
    
    # Reject if result is too short after normalization
    return result if len(result) >= 3 else ""


def _extract_id_from_text(text: str) -> str:
    """Extract student ID from text with improved pattern matching."""
    # 1. Try labeled ID patterns (highest priority)
    label_match = _ID_LABEL_RE.search(text)
    if label_match:
        extracted = label_match.group(1).strip().upper()
        # Validate ID is not just noise
        if len(extracted) >= 4 and not extracted.startswith("NO"):
            return extracted
    
    # 2. Try token patterns (student ID codes)
    token_match = _ID_TOKEN_RE.search(text)
    if token_match:
        extracted = token_match.group(1).strip().upper()
        if len(extracted) >= 5:  # Minimum reasonable length
            return extracted
    
    return ""


def _extract_name_from_text(text: str) -> str:
    """Extract student name from text with validation."""
    # 1. Try explicit label patterns (highest priority)
    label_match = _NAME_LABEL_RE.search(text)
    if label_match:
        extracted = _normalize_name(label_match.group(1))
        if extracted:
            return extracted
    
    # 2. Try pattern extraction from general text
    collapsed = re.sub(r"[_\-]+", " ", text)
    collapsed = re.sub(r"\b\d+\b", " ", collapsed)
    collapsed = _ID_TOKEN_RE.sub(" ", collapsed)
    
    # Extract words that look like names
    words = [
        w for w in re.findall(r"[A-Za-z][A-Za-z'\.]{1,}", collapsed)
        if w.lower() not in _NOISE_TOKENS and len(w) >= 2
    ]
    
    # Need at least 2 words for a reasonable name
    if len(words) < 2:
        return ""
    
    # Keep first plausible full name chunk (2-5 words)
    candidate = " ".join(words[:5])
    extracted = _normalize_name(candidate)
    
    return extracted if extracted else ""


def _infer_identity_for_group(files: list[dict], base_dir: str, lms_meta: dict) -> dict:
    """
    Infer student name/id using strict priority:
    1) Submission filename
    2) Folder names
    3) Document content
    """
    inferred_name = ""
    inferred_id = ""
    name_source = ""
    id_source = ""

    # 1) Prefer metadata in uploaded file name(s)
    for item in files:
        stem = Path(item["filename"]).stem
        file_name = _extract_name_from_text(stem)
        file_id = _extract_id_from_text(stem)
        if file_name and not inferred_name:
            inferred_name = file_name
            name_source = "filename"
        if file_id and not inferred_id:
            inferred_id = file_id
            id_source = "filename"
        if inferred_name and inferred_id:
            break

    # 2) Fall back to folder names when filename is missing metadata
    if not inferred_name or not inferred_id:
        for item in files:
            rel_parts = Path(os.path.relpath(item["path"], base_dir)).parts[:-1]
            for folder_part in rel_parts:
                folder_name = _extract_name_from_text(folder_part)
                folder_id = _extract_id_from_text(folder_part)
                if folder_name and not inferred_name:
                    inferred_name = folder_name
                    name_source = "folder"
                if folder_id and not inferred_id:
                    inferred_id = folder_id
                    id_source = "folder"
                if inferred_name and inferred_id:
                    break
            if inferred_name and inferred_id:
                break

    # 3) Last fallback: parse explicit labels from document content
    if not inferred_name or not inferred_id:
        for item in files:
            content = item.get("content", "")
            if not content or item.get("error"):
                continue
            preview = content[:15000]
            content_name = _extract_name_from_text(preview)
            content_id = _extract_id_from_text(preview)
            if content_name and not inferred_name:
                inferred_name = content_name
                name_source = "content"
            if content_id and not inferred_id:
                inferred_id = content_id
                id_source = "content"
            if inferred_name and inferred_id:
                break

    if not inferred_name and lms_meta.get("student_name"):
        inferred_name = lms_meta["student_name"]
        name_source = "lms"

    return {
        "name": inferred_name,
        "id": inferred_id,
        "name_source": name_source,
        "id_source": id_source,
    }


def _normalize_key(value: str) -> str:
    """Normalize names/keys for robust roster matching."""
    value = (value or "").strip().lower()
    value = re.sub(r"\s+", " ", value)
    value = re.sub(r"[^a-z0-9 ]", "", value)
    return value.strip()


def load_student_roster(excel_path: str) -> list[dict]:
    """
    Load a student roster from Excel and return rows as:
    [{"name": "...", "id": "..."}]

    Header matching is heuristic and accepts common variants:
    - Name: name, student name, full name
    - ID: id, student id, roll no, registration no, reg no
    """
    wb = openpyxl.load_workbook(excel_path, read_only=True, data_only=True)
    ws = wb.active

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    headers = [str(h).strip().lower() if h is not None else "" for h in rows[0]]
    name_idx = None
    id_idx = None
    for idx, header in enumerate(headers):
        if name_idx is None and header in {"name", "student name", "full name"}:
            name_idx = idx
        if id_idx is None and header in {
            "id", "student id", "roll no", "roll number", "registration no", "reg no"
        }:
            id_idx = idx

    # Fallback header detection by token presence
    if name_idx is None:
        for idx, header in enumerate(headers):
            if "name" in header:
                name_idx = idx
                break
    if id_idx is None:
        for idx, header in enumerate(headers):
            if "id" in header or "roll" in header or "reg" in header:
                id_idx = idx
                break

    if name_idx is None or id_idx is None:
        raise ValueError(
            "Roster Excel must include recognizable Name and ID columns in the first row."
        )

    roster: list[dict] = []
    for row in rows[1:]:
        if row is None:
            continue
        raw_name = row[name_idx] if name_idx < len(row) else None
        raw_id = row[id_idx] if id_idx < len(row) else None
        name = _normalize_name(str(raw_name or ""))
        sid = str(raw_id or "").strip().upper()
        if not name and not sid:
            continue
        if not name or not sid:
            continue
        roster.append({"name": name, "id": sid})
    return roster


def _build_roster_indexes(roster: list[dict]) -> tuple[dict[str, dict], dict[str, dict]]:
    """Build normalized lookup indexes for roster rows."""
    by_name: dict[str, dict] = {}
    by_id: dict[str, dict] = {}
    for row in roster:
        nkey = _normalize_key(row.get("name", ""))
        ikey = str(row.get("id", "")).strip().upper()
        if nkey:
            by_name[nkey] = row
        if ikey:
            by_id[ikey] = row
    return by_name, by_id


def _extract_candidate_tokens(files: list[dict], base_dir: str, lms_meta: dict) -> tuple[set[str], set[str]]:
    """Collect possible name/id tokens from paths only (never from document content)."""
    names: set[str] = set()
    ids: set[str] = set()

    if lms_meta.get("student_name"):
        names.add(_normalize_key(lms_meta["student_name"]))

    for item in files:
        stem = Path(item["filename"]).stem
        maybe_name = _extract_name_from_text(stem)
        maybe_id = _extract_id_from_text(stem)
        if maybe_name:
            names.add(_normalize_key(maybe_name))
        if maybe_id:
            ids.add(maybe_id.strip().upper())

        rel_parts = Path(os.path.relpath(item["path"], base_dir)).parts[:-1]
        for part in rel_parts:
            folder_name = _extract_name_from_text(part)
            folder_id = _extract_id_from_text(part)
            if folder_name:
                names.add(_normalize_key(folder_name))
            if folder_id:
                ids.add(folder_id.strip().upper())

    return names, ids


def _match_roster_identity(
    files: list[dict],
    base_dir: str,
    lms_meta: dict,
    roster_by_name: dict[str, dict],
    roster_by_id: dict[str, dict],
) -> dict:
    """
    Resolve identity from roster. IDs/names are always taken from roster rows.
    Matching only inspects path-derived hints (filename/folder/LMS metadata),
    never submission content.
    """
    candidate_names, candidate_ids = _extract_candidate_tokens(files, base_dir, lms_meta)

    # 1) Exact ID match if any ID-like token appears in path metadata.
    for cid in candidate_ids:
        if cid in roster_by_id:
            row = roster_by_id[cid]
            return {"name": row["name"], "id": row["id"], "name_source": "roster", "id_source": "roster"}

    # 2) Exact normalized name match.
    for cname in candidate_names:
        if cname in roster_by_name:
            row = roster_by_name[cname]
            return {"name": row["name"], "id": row["id"], "name_source": "roster", "id_source": "roster"}

    # 3) Containment match for slight formatting differences.
    for cname in candidate_names:
        for roster_name_key, row in roster_by_name.items():
            if cname and roster_name_key and (
                cname in roster_name_key or roster_name_key in cname
            ):
                return {"name": row["name"], "id": row["id"], "name_source": "roster", "id_source": "roster"}

    # Keep deterministic placeholders so downstream grading does not fall back
    # to LLM/content-based identity extraction when roster mode is enabled.
    return {
        "name": "Unmatched Submission",
        "id": "N/A",
        "name_source": "roster_unmatched",
        "id_source": "roster_unmatched",
    }


def read_file(file_path: str) -> str:
    """Read a single file based on its extension. Returns extracted text."""
    ext = Path(file_path).suffix.lower()
    reader = _READERS.get(ext)
    if reader is None:
        raise ValueError(f"Unsupported file format: {ext}")
    return reader(file_path)


def _parse_lms_path(file_path: str, base_dir: str) -> dict:
    """
    Extract student name and assignment metadata from LMS folder structure.

    Expected structure:
      <base>/<id>-<course>-<batch>-<section> - <semester>-<assignment>-<id>/
        <STUDENT NAME>_<number>_assignsubmission_file/
          actual_file.ext
    """
    rel = os.path.relpath(file_path, base_dir)
    parts = Path(rel).parts

    student_name = ""
    assignment_name = ""
    course_code = ""
    semester = ""

    for part in parts:
        # Student name folder: "ABDUL REHMAN NAEEM RIAZ_837796_assignsubmission_file"
        name_match = re.match(r"^([A-Za-z][A-Za-z\s]+?)_\d+_assignsubmission", part)
        if name_match:
            student_name = " ".join(name_match.group(1).split()).title()

        # Assignment name from parent folder
        assign_match = re.search(r"(Assignment\s+\d+)", part, re.IGNORECASE)
        if assign_match:
            assignment_name = assign_match.group(1).title()

        # Course code patterns: CC323, CS-101, EE 201
        course_match = re.search(
            r"\b([A-Z]{2,6}\s*-\s*\d{2,4}|[A-Z]{2,6}\s+\d{2,4}|[A-Z]{2,6}\d{2,4})\b",
            part,
        )
        if course_match:
            course_code = re.sub(r"\s+", " ", course_match.group(1).replace(" - ", "-")).strip()

        # Semester patterns: F2025, SP25, 2025-Spring, Spring-2025
        sem_match = re.search(
            r"\b((?:F|S|SP|FA)\d{2,4}|\d{4}\s*[-_/]\s*(?:spring|summer|fall|autumn)|(?:spring|summer|fall|autumn)\s*[-_/]\s*\d{4})\b",
            part,
            re.IGNORECASE,
        )
        if sem_match:
            semester = sem_match.group(1).upper()

    return {
        "student_name":    student_name,
        "assignment_name": assignment_name,
        "course_code":     course_code,
        "semester":        semester,
    }


def _extract_nested_archives(directory: str, max_passes: int = 3) -> None:
    """
    Extract nested archive files found inside the extracted LMS bundle.
    Handles "ZIP/RAR/7z within ZIP" submissions by recursively unpacking a few levels.
    """
    for _ in range(max_passes):
        extracted_any = False

        for root, _dirs, files in os.walk(directory):
            if any(part.startswith(".") or part.startswith("__") for part in Path(root).parts):
                continue

            for filename in files:
                ext = Path(filename).suffix.lower()
                if ext not in NESTED_ARCHIVE_EXTENSIONS:
                    continue

                archive_path = os.path.join(root, filename)
                target_dir = os.path.join(root, f"{Path(filename).stem}_nested_{ext.lstrip('.')}")
                if os.path.exists(target_dir):
                    continue

                try:
                    extract_archive(archive_path, target_dir)
                    extracted_any = True
                    logger.info("Extracted nested archive: %s", archive_path)
                except Exception as e:
                    logger.warning("Skipping nested archive %s: %s", archive_path, e)

        if not extracted_any:
            break


def _student_group_key(file_path: str, base_dir: str, lms_meta: dict) -> str:
    """Derive a stable grouping key so each student folder becomes one submission."""
    rel_parts = Path(os.path.relpath(file_path, base_dir)).parts

    for part in rel_parts:
        if re.match(r"^.+?_\d+_assignsubmission", part, re.IGNORECASE):
            return part

    if lms_meta.get("student_name"):
        return f"student::{lms_meta['student_name'].lower()}"

    if len(rel_parts) >= 2:
        if re.search(r"assignment|submissions?|class|course|section|semester", rel_parts[0], re.IGNORECASE):
            return rel_parts[1]
        return rel_parts[0]

    return "root_submission"


def collect_submissions(
    directory: str,
    exclude_filenames: list[str] | None = None,
    student_roster: list[dict] | None = None,
) -> list[dict]:
    """
    Walk through an extracted submissions directory and read every
    supported file.

    Parameters
    ----------
    directory         : Path to the extracted ZIP directory.
    exclude_filenames : Optional list of filenames to skip (e.g. answer key).

    Returns
    -------
    list[dict] — one combined submission per student folder.
    """
    exclude_set = {f.lower() for f in (exclude_filenames or [])}
    grouped: dict[str, dict] = {}
    roster_by_name, roster_by_id = _build_roster_indexes(student_roster or [])

    # First expand any student-provided archive uploads (ZIP inside ZIP).
    _extract_nested_archives(directory)

    from utils.cache import _make_safe_key

    for root, _dirs, files in os.walk(directory):
        # Skip hidden / system directories
        if any(
            part.startswith(".") or part.startswith("__")
            for part in Path(root).parts
        ):
            continue

        for filename in sorted(files):
            ext = Path(filename).suffix.lower()
            if ext in NESTED_ARCHIVE_EXTENSIONS:
                extracted_dir = os.path.join(root, f"{Path(filename).stem}_nested_{ext.lstrip('.')}")
                if os.path.isdir(extracted_dir):
                    continue
                # If an archive reaches this pass, nested extraction failed or no
                # local extractor is available. Keep a readable error in the report.
                full_path = os.path.join(root, filename)
                rel_path = os.path.relpath(full_path, directory)
                lms_meta = _parse_lms_path(full_path, directory)
                group_key = _student_group_key(full_path, directory, lms_meta)
                group = grouped.setdefault(group_key, {"files": [], "lms_meta": lms_meta})
                group["files"].append(
                    {
                        "filename": filename,
                        "path": full_path,
                        "rel_path": rel_path,
                        "content": (
                            f"[ERROR: Could not extract archive '{filename}'. "
                            "RAR/7z submissions require a local archive tool such as bsdtar/libarchive or 7z. "
                            "Student may need to resubmit as ZIP.]"
                        ),
                        "error": "archive_extract_failed",
                    }
                )
                continue
            if ext not in SUPPORTED_EXTENSIONS:
                continue

            # Skip answer key or any other excluded file
            if filename.lower() in exclude_set:
                logger.info("Skipping excluded file: %s", filename)
                continue

            full_path = os.path.join(root, filename)
            rel_path = os.path.relpath(full_path, directory)
            try:
                file_size_mb = os.path.getsize(full_path) / (1024 * 1024)
                if file_size_mb > MAX_FILE_SIZE_MB:
                    logger.warning(
                        "Skipping %s — file size %.1fMB exceeds limit of %dMB",
                        filename, file_size_mb, MAX_FILE_SIZE_MB,
                    )
                    content = f"[SKIPPED: File too large ({file_size_mb:.1f}MB). Student must resubmit.]"
                    error = "file_too_large"
                else:
                    content = read_file(full_path)
                    error = None
            except Exception as e:
                logger.error("Failed to read %s: %s", filename, e)
                content = "[ERROR: File could not be read. Student must resubmit.]"
                error = "read_failed"

            lms_meta = _parse_lms_path(full_path, directory)
            group_key = _student_group_key(full_path, directory, lms_meta)
            group = grouped.setdefault(
                group_key,
                {
                    "files": [],
                    "lms_meta": lms_meta,
                },
            )

            # Prefer richer LMS metadata if a later file has populated fields.
            if not group["lms_meta"].get("student_name") and lms_meta.get("student_name"):
                group["lms_meta"] = lms_meta

            group["files"].append(
                {
                    "filename": filename,
                    "path": full_path,
                    "rel_path": rel_path,
                    "content": content,
                    "error": error,
                }
            )

    submissions: list[dict] = []
    for group_key in sorted(grouped):
        group = grouped[group_key]
        files = sorted(group["files"], key=lambda f: f["rel_path"].lower())
        lms_meta = group["lms_meta"]
        if student_roster:
            identity_meta = _match_roster_identity(
                files=files,
                base_dir=directory,
                lms_meta=lms_meta,
                roster_by_name=roster_by_name,
                roster_by_id=roster_by_id,
            )
        else:
            identity_meta = _infer_identity_for_group(files, directory, lms_meta)

        combined_parts = []
        combined_error = None
        skipped_files = []
        for item in files:
            if item.get("error"):
                # Exclude placeholder errors from graded submission content.
                combined_error = combined_error or item["error"]
                skipped_files.append(item["filename"])
                continue
            combined_parts.append(
                f"\n\n===== FILE: {item['rel_path']} =====\n{item['content']}"
            )

        student_label = identity_meta.get("name") or lms_meta.get("student_name") or group_key
        combined_filename = student_label  # Clean name — no suffix needed
        combined_content = "".join(combined_parts).strip()

        # If all files failed, store a clear error message as content
        if not combined_content and (combined_error or skipped_files):
            file_list = ", ".join(skipped_files) if skipped_files else "all files"
            combined_content = (
                f"[ERROR: Could not read submission. "
                f"Failed files: {file_list}. "
                f"Likely scanned images — student must resubmit as searchable PDF.]"
            )

        cache_key = _make_safe_key(combined_filename, combined_content)

        submissions.append(
            {
                "filename":     combined_filename,
                "path":         files[0]["path"],
                "content":      combined_content,
                "cache_key":    cache_key,
                "lms_meta":     lms_meta,
                "identity_meta": identity_meta,
                "source_files": [f["path"] for f in files],
                **({"error": combined_error} if combined_error else {}),
            }
        )

    return submissions


def extract_and_collect(
    zip_path: str,
    exclude_filenames: list[str] | None = None,
    student_roster: list[dict] | None = None,
) -> tuple[list[dict], str]:
    """
    Extract a ZIP file and collect all supported submissions.

    Returns (submissions, extract_dir) — caller is responsible for
    cleaning up extract_dir after grading completes. This preserves
    the cache file for crash recovery during long grading sessions.

    Parameters
    ----------
    zip_path          : Path to the ZIP file.
    exclude_filenames : Optional filenames to exclude (e.g. answer key filename).
    """
    extract_dir = extract_zip(zip_path)
    submissions = collect_submissions(
        extract_dir,
        exclude_filenames=exclude_filenames,
        student_roster=student_roster,
    )
    return submissions, extract_dir
